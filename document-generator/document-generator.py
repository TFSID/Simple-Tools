from docx import Document
from fpdf import FPDF

# Content to export
content = """
# SOC Tier 3 Success Framework

## 1. Mindset & Culture
- Responsibility Mindset
- Curiosity & Continuous Learning
- Analytical Thinking
- Teamwork & Leadership
- Stress Management

## 2. Technical Mastery
- Forensic & Malware Analysis
- Threat Hunting & Detection Engineering
- SIEM/EDR Tuning & Optimization
- Threat Intelligence Integration
- Incident Response Mastery

## 3. Operational Routine
### Daily, Weekly, Monthly, Quarterly Checklists

## 4. Collaboration & Process
## 5. Key Metrics for Success
## 6. Development Plan
## 7. Tools Mastery
## 8. Shift Readiness
## 9. Exit Readiness
"""

# DOCX Generator
doc = Document()
for line in content.split('\n'):
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.strip():
        doc.add_paragraph(line)
doc.save("SOC_Tier3_Success_Framework.docx")

# PDF Generator
pdf = FPDF()
pdf.add_page()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.set_font("Arial", size=12)
for line in content.split('\n'):
    pdf.multi_cell(0, 10, txt=line, align='L')
pdf.output("SOC_Tier3_Success_Framework.pdf")

print("Files generated: DOCX and PDF")
